"""
Export Word (.docx) « Fiche plan de gestion » — partie enjeux et FCR.

Produit un document rédigé calqué sur le modèle « Modèle_Export_CICADA_PG » :

  LES ENJEUX ET FACTEURS CLÉS DE RÉUSSITE
    Les enjeux écologiques
      <intitulé enjeu> – priorité X
        <état de l'enjeu> / <détails>
        Tableau : Espèces / Habitats / Patrimoine géologique
        <OLT> → <niveaux d'exigence>
        Facteurs d'influence → pressions
    Les enjeux socio-économiques et paysagers
      …
    Les facteurs clés de réussite
      <FCR> – priorité X …

Point d'entrée public : :func:`build_plan_docx`.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

from . import export_theme

_PRIMARY = RGBColor(0x02, 0x53, 0x59)


def _appliquer_couleur_instance():
    """Réaligne la couleur des titres sur celle de l'instance (#601)."""
    global _PRIMARY

    _PRIMARY = RGBColor(*export_theme.rgb())


_TERRA = RGBColor(0xB7, 0x4D, 0x5D)
_GRAY = RGBColor(0x74, 0x6F, 0x6E)

# Flags booléens marquant un enjeu écologique
_ECO_FLAGS = (
    "categorie_ecologique", "habitat", "espece", "patrimoine_geologique",
    "fonctionnalite_ecosysteme", "autre_ecologique",
)


def _txt(value) -> str:
    if value is None:
        return ""
    return value.strip() if isinstance(value, str) else str(value)


def _is_fcr(enjeu) -> bool:
    return bool(enjeu.id_categorie and enjeu.id_categorie.mnemonique == "FCR")


def _is_ecologique(enjeu) -> bool:
    return any(getattr(enjeu, f, False) for f in _ECO_FLAGS)


def _priorite_label(enjeu) -> str:
    imp = getattr(enjeu, "id_importance", None)
    return _txt(getattr(imp, "label", "")) if imp else ""


# ---------------------------------------------------------------------------
# Helpers de mise en forme
# ---------------------------------------------------------------------------

def _heading(doc, text, level, color=None):
    # Résolu à l'appel : une valeur par défaut serait figée à l'import et
    # ignorerait la couleur d'instance (#601).
    color = color or _PRIMARY
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    return h


def _enjeu_title(doc, name, priorite):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = _PRIMARY
    if priorite:
        sep = p.add_run("  –  ")
        sep.font.color.rgb = _GRAY
        pr = p.add_run(priorite)
        pr.bold = True
        pr.font.color.rgb = _TERRA
    return p


def _sub_title(doc, text, color=None):
    """Titre intermédiaire (type OLT)."""
    color = color or _PRIMARY
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = color
    return p


def _body(doc, text, *, italic=False, gray=False):
    text = _txt(text)
    if not text:
        return None
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    if gray:
        run.font.color.rgb = _GRAY
        run.font.size = Pt(9.5)
    return p


def _bullet(doc, text, *, level=0, bold=False):
    text = _txt(text)
    if not text:
        return None
    style = "List Bullet" if level == 0 else f"List Bullet {min(level + 1, 3)}"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.bold = bold
    return p


def _label_value(doc, label, value):
    """Paragraphe « Label : valeur » (valeur en clair)."""
    value = _txt(value)
    if not value:
        return None
    p = doc.add_paragraph()
    lab = p.add_run(f"{label} : ")
    lab.bold = True
    lab.font.color.rgb = _GRAY
    p.add_run(value)
    return p


# ---------------------------------------------------------------------------
# Tableau espèces / habitats / patrimoine géologique
# ---------------------------------------------------------------------------

def _habitat_codes(enjeux) -> dict:
    """cd_hab (str) → code de l'habitat dans sa typologie (`lb_code`), #628.

    Une seule requête HabRef pour tout le plan. Renvoie un dict vide si le
    référentiel n'est pas chargé (dev/tests) ou si aucun habitat n'a de code.
    """
    from apps.habitats.models import Habref

    ints = set()
    for enjeu in enjeux:
        for hab in enjeu.habitats.all():
            raw = _txt(getattr(hab, "cd_hab", ""))
            if raw.isdigit():
                ints.add(int(raw))
    if not ints:
        return {}
    return {
        str(cd): _txt(code)
        for cd, code in Habref.objects.filter(cd_hab__in=ints).values_list("cd_hab", "lb_code")
        if _txt(code)
    }


def _bio_lists(enjeu, hab_codes=None):
    hab_codes = hab_codes or {}
    especes = [
        (_txt(t.nom_complet) + (f" ({_txt(t.nom_vern)})" if _txt(getattr(t, "nom_vern", "")) else "")).strip()
        for t in enjeu.taxons.all()
    ]
    habitats = []
    for h in enjeu.habitats.all():
        cd_hab = _txt(getattr(h, "cd_hab", ""))
        label = _txt(getattr(h, "lb_hab_fr", "")) or cd_hab
        # #628 — le code affiché est celui de la typologie (lb_code), pas le cd_hab HabRef
        code = hab_codes.get(cd_hab, "")
        habitats.append(f"{label} ({code})" if code and label else label)
    geol = [_txt(g.nom) or _txt(g.id_inpg) for g in enjeu.geologies.all()]
    for og in enjeu.objets_geologiques.all():
        lbl = _txt(getattr(og.id_objet_geologique, "label", ""))
        prec = _txt(getattr(og, "precision", ""))
        geol.append(" — ".join(x for x in (lbl, prec) if x))
    especes = [e for e in especes if e]
    habitats = [h for h in habitats if h]
    geol = [g for g in geol if g]
    return especes, habitats, geol


def _add_bio_table(doc, enjeu, table_index, enjeu_name, hab_codes=None):
    especes, habitats, geol = _bio_lists(enjeu, hab_codes)
    if not (especes or habitats or geol):
        return
    cap = doc.add_paragraph()
    cr = cap.add_run(
        f"Tableau {table_index} : Espèces, habitats et patrimoine géologique "
        f"à responsabilité dans l'enjeu « {enjeu_name} »"
    )
    cr.italic = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = _GRAY

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(table.rows[0].cells, ("Espèces", "Habitats", "Patrimoine géologique")):
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        run.font.color.rgb = _PRIMARY

    n = max(len(especes), len(habitats), len(geol))
    for i in range(n):
        row = table.add_row().cells
        row[0].text = especes[i] if i < len(especes) else ""
        row[1].text = habitats[i] if i < len(habitats) else ""
        row[2].text = geol[i] if i < len(geol) else ""


# ---------------------------------------------------------------------------
# Rendu d'un enjeu / FCR
# ---------------------------------------------------------------------------

def _manual_rank_map(items, id_getter):
    """id → numéro d'affichage, réplique de la numérotation du frontend (#442/#526).

    Un ``numero_manuel`` fixé est réservé ; l'auto-numérotation des autres saute
    les indices occupés. ``items`` doit être dans l'ordre d'affichage voulu.
    """
    reserved = {it.numero_manuel for it in items if getattr(it, "numero_manuel", None)}
    ranks = {}
    auto = 0
    for it in items:
        manuel = getattr(it, "numero_manuel", None)
        if manuel:
            ranks[id_getter(it)] = manuel
        else:
            auto += 1
            while auto in reserved:
                auto += 1
            ranks[id_getter(it)] = auto
    return ranks


def _render_enjeu(doc, enjeu, table_counter, *, is_fcr=False, hab_codes=None,
                  enjeu_rank=None, olt_ranks=None):
    olt_ranks = olt_ranks or {}
    name = _txt(enjeu.libelle) or _txt(enjeu.intitule_court) or f"Enjeu {enjeu.id_enjeu}"
    # #628 — préfixer par le numéro d'affichage (« Enjeu 1 : … »), hors FCR
    if not is_fcr and enjeu_rank:
        name = f"Enjeu {enjeu_rank} : {name}"
    _enjeu_title(doc, name, _priorite_label(enjeu))
    _label_value(doc, "Précision sur l'état de l'enjeu", enjeu.etat_enjeu)
    _body(doc, enjeu.description)

    if not is_fcr:
        table_counter[0] += 1
        _add_bio_table(doc, enjeu, table_counter[0], name, hab_codes)

    # Objectifs à long terme → niveaux d'exigence
    for olt in enjeu.objectifs_long_terme.all():
        # #628 — préfixer par le numéro global de l'OLT (« OLT 1 : … »)
        olt_rank = olt_ranks.get(olt.id_olt)
        olt_prefix = f"OLT {olt_rank} : " if olt_rank else "OLT : "
        _sub_title(doc, f"{olt_prefix}{_txt(olt.libelle)}")
        _body(doc, olt.description)
        niveaux = list(olt.niveaux_exigence.all())
        if niveaux:
            # #628 — intertitre explicite, comme pour « Facteurs d'influence »
            _sub_title(doc, "Niveaux d'exigence", color=_TERRA)
        for ne in niveaux:
            _bullet(doc, ne.libelle)
            _body(doc, ne.description, gray=True)

    # Facteurs d'influence → pressions
    facteurs = list(enjeu.facteurs_influence.all())
    if facteurs:
        _sub_title(doc, "Facteurs d'influence", color=_TERRA)
        for facteur in facteurs:
            _bullet(doc, facteur.libelle, bold=True)
            _body(doc, facteur.description, gray=True)
            for pression in facteur.pressions.all():
                _bullet(doc, pression.libelle, level=1)

    # OO rattachés directement (surtout FCR)
    from .models_enjeux import ObjectifOperationnel
    direct_oo = ObjectifOperationnel.objects.filter(id_enjeu=enjeu).prefetch_related(
        "resultats_attendus"
    )
    if direct_oo:
        _sub_title(doc, "Objectifs opérationnels", color=_TERRA)
        for oo in direct_oo:
            _bullet(doc, oo.libelle, bold=True)
            _body(doc, oo.description, gray=True)
            for ra in oo.resultats_attendus.all():
                _bullet(doc, ra.libelle, level=1)


# ---------------------------------------------------------------------------
# Point d'entrée
# ---------------------------------------------------------------------------

def _prefetched_enjeux(plan):
    from .models_enjeux import Enjeu
    qs = (
        Enjeu.objects.filter(id_pg=plan)
        .select_related("id_categorie", "id_categorie_fcr", "id_importance")
        .prefetch_related(
            "taxons", "habitats", "geologies", "objets_geologiques__id_objet_geologique",
            "objectifs_long_terme__niveaux_exigence",
            "facteurs_influence__pressions",
        )
        .order_by("ordre", "id_enjeu")
    )
    return list(qs)


def build_plan_docx(plan) -> bytes:
    """Construit la fiche Word (enjeux + FCR) du plan de gestion."""
    _appliquer_couleur_instance()
    doc = Document()
    table_counter = [0]

    _heading(doc, "LES ENJEUX ET FACTEURS CLÉS DE RÉUSSITE", 1)

    enjeux = _prefetched_enjeux(plan)
    hab_codes = _habitat_codes(enjeux)
    ecologiques = [e for e in enjeux if not _is_fcr(e) and _is_ecologique(e)]
    socioeco = [e for e in enjeux if not _is_fcr(e) and not _is_ecologique(e)]
    fcrs = [e for e in enjeux if _is_fcr(e)]

    # #628 — numéros d'affichage. Enjeux : sur la liste non-FCR ordonnée (comme
    # le frontend). OLT : numérotation globale continue (enjeux puis FCR).
    non_fcr = [e for e in enjeux if not _is_fcr(e)]
    enjeu_ranks = _manual_rank_map(non_fcr, lambda e: e.id_enjeu)
    olt_items = [olt for e in non_fcr for olt in e.objectifs_long_terme.all()]
    olt_items += [olt for e in fcrs for olt in e.objectifs_long_terme.all()]
    olt_ranks = _manual_rank_map(olt_items, lambda o: o.id_olt)

    _heading(doc, "Les enjeux écologiques", 2)
    if ecologiques:
        for enjeu in ecologiques:
            _render_enjeu(doc, enjeu, table_counter, hab_codes=hab_codes,
                          enjeu_rank=enjeu_ranks.get(enjeu.id_enjeu), olt_ranks=olt_ranks)
    else:
        _body(doc, "Aucun enjeu écologique renseigné.", italic=True, gray=True)

    _heading(doc, "Les enjeux socio-économiques et paysagers", 2)
    if socioeco:
        for enjeu in socioeco:
            _render_enjeu(doc, enjeu, table_counter, hab_codes=hab_codes,
                          enjeu_rank=enjeu_ranks.get(enjeu.id_enjeu), olt_ranks=olt_ranks)
    else:
        _body(doc, "Aucun enjeu socio-économique ou paysager renseigné.", italic=True, gray=True)

    _heading(doc, "Les facteurs clés de réussite", 2)
    if fcrs:
        for fcr in fcrs:
            _render_enjeu(doc, fcr, table_counter, is_fcr=True, hab_codes=hab_codes,
                          olt_ranks=olt_ranks)
    else:
        _body(doc, "Aucun facteur clé de réussite renseigné.", italic=True, gray=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
